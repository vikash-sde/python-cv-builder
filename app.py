from docx import Document
from docx.shared import Inches

document = Document()

# profile picture
document.add_picture('profile.png',width=Inches(1.25))

# name phone number and email details
name = input('name? ')
phone_number = input('phone number? ')
email = input('email? ')

document.add_paragraph(
    name + '|' + phone_number + '|' + email
)

# about me

document.add_heading('About me')
document.add_paragraph(input('About yourself? '))

# work experience
document.add_heading('Work experience ')
p = document.add_paragraph()

company = input('Company ')
from_date = input('From Date ')
to_date = input('To Date ')

p.add_run(company + ' - ').bold = True
p.add_run(from_date + ' - ' + to_date + '\n').italic = True

experience_details = input('Experience Details ' + company)
p.add_run(experience_details)

# more experiences
while True:
    has_more_experiences = input('Do you have more experiences? (yes/no): ').lower()
    if has_more_experiences == 'yes':
        p = document.add_paragraph()

        company = input('Company ')
        from_date = input('From Date ')
        to_date = input('To Date ')

        p.add_run(company + ' - ').bold = True
        p.add_run(from_date + ' - ' + to_date + '\n').italic = True

        experience_details = input('Experience Details ' + company)
        p.add_run(experience_details)
    else:
        break


document.save('cv.docx')